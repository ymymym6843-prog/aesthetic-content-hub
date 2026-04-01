"""
Extract skin MBTI data from docx to JS data file.
Source: 스킨_MBTI_책자_완성본 part 2.docx
Output: skin_mbti_booklet/data/skin_types_data.js
"""
import sys
import io
import re
import json
import os
from datetime import datetime

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import docx

DOCX_PATH = "C:/Users/dbals/AppData/Local/Temp/skin_mbti.docx"
OUTPUT_PATH = "C:/Users/dbals/VibeCoding/IM_AESTHETICS/skin_mbti_booklet/data/skin_types_data.js"

TYPE_CODES = [
    "OSNT", "OSNW", "OSPT", "OSPW",
    "ORNT", "ORNW", "ORPT", "ORPW",
    "DSNT", "DSNW", "DSPT", "DSPW",
    "DRNT", "DRNW", "DRPT", "DRPW",
]

# Indicator mappings derived from code letters
INDICATOR_MAP = {
    "O": "moisture", "D": "moisture",
    "S": "sensitivity", "R": "sensitivity",
    "N": "pigment", "P": "pigment",
    "T": "aging", "W": "aging",
}

def get_table_text(table):
    """Return all text from a table, cells joined by newlines."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    return rows


def get_cell_text(cell):
    return cell.text.strip()


def get_table_first_cell(table):
    if table.rows and table.rows[0].cells:
        return table.rows[0].cells[0].text.strip()
    return ""


def build_element_sequence(doc):
    """Build ordered list of (type, index, element) for all body elements."""
    body = doc.element.body
    sequence = []
    p_idx = 0
    t_idx = 0
    for elem in body:
        tag = elem.tag.split("}")[1] if "}" in elem.tag else elem.tag
        if tag == "p":
            sequence.append(("p", p_idx, elem))
            p_idx += 1
        elif tag == "tbl":
            sequence.append(("t", t_idx, elem))
            t_idx += 1
    return sequence


def para_text(elem):
    """Extract clean text from a paragraph XML element."""
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    parts = []
    for t in elem.iter(f"{{{ns}}}t"):
        if t.text:
            parts.append(t.text)
    return "".join(parts)


def extract_ch01_indicators(doc):
    """Extract 4 indicator pairs from Ch.01 tables (7-14)."""
    indicators = []
    # Tables 7-14 in pairs: (7,8), (9,10), (11,12), (13,14)
    pairs = [(7, 8), (9, 10), (11, 12), (13, 14)]
    for header_idx, detail_idx in pairs:
        header_table = doc.tables[header_idx]
        detail_table = doc.tables[detail_idx]
        # Header table: [pair_code, title+subtitle, indicator_num]
        header_row = header_table.rows[0].cells
        pair_code = header_row[0].text.strip()  # e.g. "O vs D"
        title_raw = header_row[1].text.strip()  # "유수분 밸런스\n\"...\""
        title_lines = title_raw.split("\n")
        title = title_lines[0].strip()
        subtitle = title_lines[1].strip().strip('"') if len(title_lines) > 1 else ""

        # Detail table: [left_desc, right_desc]
        detail_row = detail_table.rows[0].cells
        left_raw = detail_row[0].text.strip()
        right_raw = detail_row[1].text.strip()

        def parse_indicator(raw):
            lines = raw.split("\n")
            # First line: "[X] Name (KR name)"
            first = lines[0].strip()
            # Extract code like [O], [D], etc.
            code_match = re.search(r"\[(\w)\]", first)
            code = code_match.group(1) if code_match else ""
            # Extract name
            name_match = re.search(r"\]\s+(.+?)$", first)
            name = name_match.group(1).strip() if name_match else first
            desc = "\n".join(l.strip().strip('"') for l in lines[1:] if l.strip())
            return {"code": code, "name": name, "description": desc}

        indicators.append({
            "pair": pair_code,
            "title": title,
            "subtitle": subtitle,
            "left": parse_indicator(left_raw),
            "right": parse_indicator(right_raw),
        })

    return indicators


def extract_ch02_solutions(doc):
    """Extract Ch.02 solutions from Table 17."""
    table = doc.tables[17]
    solutions = []
    for r, row in enumerate(table.rows):
        if r == 0:
            continue  # Skip header row
        cells = row.cells
        if len(cells) < 4:
            continue
        indicator = cells[0].text.strip()
        goal_raw = cells[1].text.strip()
        equipment = cells[2].text.strip()
        tip = cells[3].text.strip()
        # goal may have title and description separated by \n
        goal_lines = goal_raw.split("\n")
        goal_title = goal_lines[0].strip()
        goal_desc = "\n".join(l.strip().strip('"') for l in goal_lines[1:] if l.strip())
        solutions.append({
            "indicator": indicator,
            "goal": goal_title,
            "goalDescription": goal_desc,
            "equipment": equipment,
            "tip": tip,
        })
    return solutions


def extract_ch03_overview(doc):
    """Extract Ch.03 overview cards from Tables 19-26."""
    overview = []
    # Tables 19-26 each have 1 row x 3 cols (code1, empty, code2)
    for tidx in range(19, 27):
        table = doc.tables[tidx]
        row = table.rows[0].cells
        for col_idx in [0, 2]:
            raw = row[col_idx].text.strip()
            if not raw:
                continue
            lines = raw.split("\n")
            code = ""
            nickname = ""
            summary = ""
            arrow = ""
            for i, line in enumerate(lines):
                line = line.strip()
                if re.match(r"^[A-Z]{4}$", line):
                    code = line
                elif line.startswith("#"):
                    nickname = line
                elif line.startswith("→"):
                    arrow = line[1:].strip()
                elif line and not code:
                    pass  # Skip pre-code lines
                elif code and not nickname and line.startswith("#"):
                    nickname = line
                elif code and nickname and not arrow and not line.startswith("→"):
                    summary = line if not summary else summary + " " + line
            if code:
                overview.append({
                    "code": code,
                    "nickname": nickname,
                    "summary": summary,
                    "arrow": arrow,
                })
    return overview


def extract_ch04_roadmap(doc):
    """Extract Ch.04 3-phase roadmap from Tables 29, 30, 31."""
    steps = {}
    for tidx, step_key in [(29, "step1"), (30, "step2"), (31, "step3")]:
        table = doc.tables[tidx]
        row = table.rows[0].cells
        # cell[0] = "STEP\n1"
        # cell[1] = "Title — Subtitle\n목표: ...\n장비: ...\n고객 경험: ..."
        content = row[1].text.strip() if len(row) > 1 else ""
        lines = content.split("\n")
        title = ""
        subtitle = ""
        goal = ""
        equipment = ""
        customer_exp = ""
        if lines:
            first = lines[0].strip()
            if " — " in first:
                parts = first.split(" — ", 1)
                title = parts[0].strip()
                subtitle = parts[1].strip()
            else:
                title = first
        for line in lines[1:]:
            line = line.strip()
            if line.startswith("목표:"):
                goal = line[3:].strip()
            elif line.startswith("장비:"):
                equipment = line[3:].strip()
            elif line.startswith("고객 경험:"):
                customer_exp = line[6:].strip().strip('"')
        steps[step_key] = {
            "title": title,
            "subtitle": subtitle,
            "goal": goal,
            "equipment": equipment,
            "customerExperience": customer_exp,
        }
    return steps


def extract_type_overview_card(doc, type_code):
    """Extract type card from TIPY tables (Tables 33, 35, 38, 40, 43, ...)."""
    # Map type code to table index
    tipy_table_map = {
        "OSNT": 33, "OSNW": 35, "OSPT": 38, "OSPW": 40,
        "ORNT": 43, "ORNW": 45, "ORPT": 48, "ORPW": 50,
        "DSNT": 53, "DSNW": 55, "DSPT": 58, "DSPW": 60,
        "DRNT": 63, "DRNW": 65, "DRPT": 68, "DRPW": 70,
    }
    tidx = tipy_table_map.get(type_code)
    if tidx is None:
        return {}, ""
    table = doc.tables[tidx]
    row = table.rows[0].cells
    # cell[0]: "TIPY N\nCODE\n#nickname"
    # cell[1]: "Name description line\nFull characteristics..."
    left = row[0].text.strip()
    right = row[1].text.strip()

    # Parse left cell
    nickname = ""
    for line in left.split("\n"):
        line = line.strip()
        if line.startswith("#"):
            nickname = line
            break

    # Parse right cell
    right_lines = right.split("\n")
    name = right_lines[0].strip() if right_lines else ""
    characteristics = "\n".join(l.strip() for l in right_lines[1:] if l.strip())

    return {"nickname": nickname, "name": name}, characteristics


def extract_ch05_step_cards(doc, type_code):
    """Extract STEP 1/2/3 cards from Ch.05 overview tables."""
    # Map type to the step-card table index
    step_table_map = {
        "OSNT": 34, "OSNW": 36, "OSPT": 39, "OSPW": 41,
        "ORNT": 44, "ORNW": 46, "ORPT": 49, "ORPW": 51,
        "DSNT": 54, "DSNW": 56, "DSPT": 59, "DSPW": 61,
        "DRNT": 64, "DRNW": 66, "DRPT": 69, "DRPW": 71,
    }
    tidx = step_table_map.get(type_code)
    if tidx is None:
        return {}
    table = doc.tables[tidx]
    row = table.rows[0].cells
    if len(row) < 3:
        return {}

    def parse_step_cell(cell_text):
        lines = cell_text.strip().split("\n")
        title = ""
        desc = ""
        for i, line in enumerate(lines):
            line = line.strip()
            if re.match(r"^STEP\s+\d+$", line) or re.match(r"^STEP\s+undefined$", line):
                continue
            if not title:
                title = line
            else:
                desc = line if not desc else desc + " " + line
        return {"title": title, "description": desc}

    return {
        "step1": parse_step_cell(row[0].text),
        "step2": parse_step_cell(row[1].text),
        "step3": parse_step_cell(row[2].text),
    }


def extract_management_manual_steps(doc, type_code):
    """Extract the 3-step management program from the detailed manual tables."""
    # Map type to the 3-step table in the Part 2 manual section
    # These tables have 3 cols: STEP1 title+desc | STEP2 | STEP3
    # Located after the management manual header table
    manual_step_table_map = {
        "OSNT": 76, "OSNW": 88, "DSNT": 100, "DRPW": 112,
        "OSPT": 124, "OSPW": 136, "ORNT": 148,
        "ORNW": 160, "ORPT": 172, "ORPW": 184,
        "DSNW": 196, "DSPT": 208, "DSPW": 220,
        "DRNT": 232, "DRNW": 244, "DRPT": 256,
    }
    # For types we have step tables from (try to find dynamically)
    # The step tables follow pattern: after manual header, before in-care program table
    tidx = manual_step_table_map.get(type_code)
    if tidx is None or tidx >= len(doc.tables):
        return {}
    table = doc.tables[tidx]
    if len(table.rows) == 0 or len(table.rows[0].cells) < 3:
        return {}

    def parse_step_cell(cell_text):
        lines = cell_text.strip().split("\n")
        title = ""
        desc = ""
        for line in lines:
            line = line.strip()
            if re.match(r"^STEP\s+(\d+|undefined)$", line):
                continue
            if not title:
                title = line
            else:
                desc = line if not desc else desc + " " + line
        return {"title": title, "description": desc}

    return {
        "step1": parse_step_cell(table.rows[0].cells[0].text),
        "step2": parse_step_cell(table.rows[0].cells[1].text),
        "step3": parse_step_cell(table.rows[0].cells[2].text),
    }


def extract_in_care_program(doc, type_code):
    """Extract in-care program table (7-row table after 3-step cards)."""
    # These 7-row x 3-col tables: stage | content | expert tip
    incare_table_map = {
        "OSNT": 77, "OSNW": 89, "DSNT": 101, "DRPW": 113,
        "OSPT": 125, "OSPW": 137, "ORNT": 149,
        "ORNW": 161, "ORPT": 173, "ORPW": 185,
        "DSNW": 197, "DSPT": 209, "DSPW": 221,
        "DRNT": 233, "DRNW": 245, "DRPT": 257,
    }
    tidx = incare_table_map.get(type_code)
    if tidx is None or tidx >= len(doc.tables):
        return []
    table = doc.tables[tidx]
    program = []
    for r, row in enumerate(table.rows):
        cells = row.cells
        if len(cells) < 2:
            continue
        stage = cells[0].text.strip()
        content = cells[1].text.strip()
        tip = cells[2].text.strip() if len(cells) > 2 else ""
        if stage and stage != "단계":
            program.append({"stage": stage, "content": content, "tip": tip})
    return program


def extract_homecare_table(doc, type_code):
    """Extract homecare prescription from 4-row x 2-col table."""
    homecare_table_map = {
        "OSNT": 78, "OSNW": 90, "DSNT": 102, "DRPW": 114,
        "OSPT": 126, "OSPW": 138, "ORNT": 150,
        "ORNW": 162, "ORPT": 174, "ORPW": 186,
        "DSNW": 198, "DSPT": 210, "DSPW": 222,
        "DRNT": 234, "DRNW": 246, "DRPT": 258,
    }
    tidx = homecare_table_map.get(type_code)
    if tidx is None or tidx >= len(doc.tables):
        return {}
    table = doc.tables[tidx]
    homecare = {}
    key_map = {
        "세안": "cleansing",
        "클렌징": "cleansing",
        "보습": "moisturizing",
        "오일": "oil",
        "금지": "avoid",
        "스페셜": "special",
        "스팟": "spot",
        "필링": "peeling",
        "추천": "recommended",
    }
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        key_raw = cells[0].text.strip()
        val = cells[1].text.strip()
        # Map key
        mapped_key = None
        for k, v in key_map.items():
            if k in key_raw:
                mapped_key = v
                break
        if not mapped_key:
            mapped_key = key_raw.lower().replace(" ", "_") if key_raw else "other"
        homecare[mapped_key] = val
    return homecare


def extract_hooking_script(doc, type_code):
    """Extract hooking script from the 1-row x 1-col table."""
    hook_table_map = {
        "OSNT": 79, "OSNW": 91, "DSNT": 103, "DRPW": 115,
        "OSPT": 127, "OSPW": 139, "ORNT": 151,
        "ORNW": 163, "ORPT": 175, "ORPW": 187,
        "DSNW": 199, "DSPT": 211, "DSPW": 223,
        "DRNT": 235, "DRNW": 247, "DRPT": 259,
    }
    tidx = hook_table_map.get(type_code)
    if tidx is None or tidx >= len(doc.tables):
        return ""
    table = doc.tables[tidx]
    if table.rows and table.rows[0].cells:
        return table.rows[0].cells[0].text.strip().strip('"')
    return ""


def extract_diagnostic_checklist(doc, type_code):
    """Extract diagnostic checklist from paragraphs."""
    # The checklist is in paragraphs, not tables
    # Find the paragraphs that contain □ items for this type
    # The 16 type sections appear in fixed paragraph ranges, grouped by TYPE 특징 markers

    # Map type to paragraph index of its "TYPE 특징" marker (Ch05 section)
    type_para_map = {
        "OSNT": 206, "OSNW": 246, "DSNT": 285, "DSNW": 325,
        "OSPT": 365, "OSPW": 404, "DSPT": 444, "DSPW": 484,
        "ORNT": 524, "ORNW": 564, "DRNT": 604, "DRNW": 644,
        "ORPT": 684, "ORPW": 723, "DRPT": 763, "DRPW": 803,
    }

    start_para = type_para_map.get(type_code)
    if start_para is None:
        return {}

    # Find next TYPE 특징 after this one
    paras = doc.paragraphs
    next_start = len(paras)
    for code, pidx in type_para_map.items():
        if pidx > start_para and pidx < next_start:
            next_start = pidx

    # The checklist indicator tables between start_para and next_start tell us what
    # indicator each group of □ items belongs to. We need to track which table [O]/[S]/etc
    # precedes each group of items.

    # Build element sequence for the range
    body = doc.element.body
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    sequence = []
    p_idx = 0
    t_idx = 0
    for elem in body:
        tag = elem.tag.split("}")[1] if "}" in elem.tag else elem.tag
        if tag == "p":
            sequence.append(("p", p_idx, elem))
            p_idx += 1
        elif tag == "tbl":
            sequence.append(("t", t_idx, elem))
            t_idx += 1

    # Find sequence range for this type
    seq_start = None
    seq_end = None
    for i, (etype, eidx, elem) in enumerate(sequence):
        if etype == "p" and eidx == start_para:
            seq_start = i
        if etype == "p" and eidx == next_start:
            seq_end = i
            break
    if seq_start is None:
        return {}
    if seq_end is None:
        seq_end = len(sequence)

    # Walk through this section extracting checklist items
    checklist = {}
    current_indicator = None
    # Indicator labels in small header tables [O], [S], [P]/[N], [T]/[W]
    indicator_label_map = {
        "[O]": "oiliness", "[D]": "dryness",
        "[S]": "sensitivity", "[R]": "resistance",
        "[P]": "pigment", "[N]": "pigment",
        "[T]": "aging", "[W]": "aging",
    }

    for etype, eidx, elem in sequence[seq_start:seq_end]:
        if etype == "t":
            # Check if this is an indicator header table like [O], [S]
            table = doc.tables[eidx]
            if table.rows and len(table.rows[0].cells) >= 1:
                first_cell = table.rows[0].cells[0].text.strip()
                if first_cell in indicator_label_map:
                    current_indicator = indicator_label_map[first_cell]
                    if current_indicator not in checklist:
                        checklist[current_indicator] = []
        elif etype == "p":
            text = para_text(elem).strip()
            if text.startswith("□"):
                item = text[1:].strip()
                if current_indicator and item:
                    checklist[current_indicator].append(item)

    return checklist


def get_paragraph_range_text(doc, start_para_idx, end_para_idx):
    """Get text from paragraphs in a range, excluding empty."""
    texts = []
    for i in range(start_para_idx, min(end_para_idx, len(doc.paragraphs))):
        t = doc.paragraphs[i].text.strip()
        if t:
            texts.append(t)
    return texts


def extract_type_section_paragraphs(doc, type_code):
    """
    Extract paragraph-based content for a type from Ch05 section.
    Returns: characteristics, coreStrategy, inCareDesc, homecareDesc, hookingScript, consultationPoints
    """
    type_para_map = {
        "OSNT": 206, "OSNW": 246, "DSNT": 285, "DSNW": 325,
        "OSPT": 365, "OSPW": 404, "DSPT": 444, "DSPW": 484,
        "ORNT": 524, "ORNW": 564, "DRNT": 604, "DRNW": 644,
        "ORPT": 684, "ORPW": 723, "DRPT": 763, "DRPW": 803,
    }
    start = type_para_map.get(type_code)
    if start is None:
        return {}

    paras = doc.paragraphs
    # Find section boundaries
    next_starts = sorted([v for v in type_para_map.values() if v > start])
    end = next_starts[0] if next_starts else len(paras)

    result = {
        "characteristics": "",
        "coreStrategy": "",
        "inCareDesc": "",
        "homecareDesc": "",
        "hookingScript": "",
        "consultationPoints": [],
    }

    current_section = None
    i = start
    while i < end:
        p = paras[i]
        text = p.text.strip()

        if text == "TYPE 특징":
            current_section = "characteristics"
        elif text == "1. 관리 핵심 전략":
            current_section = "coreStrategy"
        elif text == "2. 인-케어 프로그램 (추천 코스)":
            current_section = "inCare"
        elif text == "3. 홈케어 처방전":
            current_section = "homecare"
        elif text == "4. 고객 상담 후킹 멘트":
            current_section = "hooking"
        elif text == "💡 원장님을 위한 상담 포인트":
            current_section = "consultation"
        elif text and current_section:
            if current_section == "characteristics" and not result["characteristics"]:
                result["characteristics"] = text
            elif current_section == "coreStrategy" and not result["coreStrategy"]:
                result["coreStrategy"] = text
            elif current_section == "inCare" and not result["inCareDesc"]:
                result["inCareDesc"] = text
            elif current_section == "homecare" and not result["homecareDesc"]:
                result["homecareDesc"] = text
            elif current_section == "hooking" and not result["hookingScript"]:
                result["hookingScript"] = text
            elif current_section == "consultation":
                if text.startswith("→"):
                    result["consultationPoints"].append(text[1:].strip())
                elif not text.startswith("다음 항목") and not text.startswith("□") and not text.startswith("["):
                    result["consultationPoints"].append(text)
        i += 1

    return result


def extract_homecare_bullets(doc, type_code):
    """Extract homecare bullet points from paragraphs (• 클렌징, • 오일, etc.)."""
    # These are in Ch05 Type overview section (Tables 33-72 section)
    # The bullets are in paragraphs after TIPY type tables but before next TIPY
    # Map type to paragraph index of homecare bullets
    homecare_para_map = {
        "OSNT": (68, 72), "OSNW": (76, 80),
        "OSPT": (85, 89), "OSPW": (93, 97),
        "ORNT": (102, 106), "ORNW": (110, 114),
        "ORPT": (119, 123), "ORPW": (127, 131),
        "DSNT": None, "DSNW": None,
        "DSPT": None, "DSPW": None,
        "DRNT": None, "DRNW": None,
        "DRPT": None, "DRPW": None,
    }
    # For DSNT onwards, the bullets may not exist in paragraphs - they're in tables
    # Let's scan dynamically using body elements

    result = {"cleansing": "", "oil": "", "recommended": "", "scalp": ""}

    range_info = homecare_para_map.get(type_code)
    if range_info is None:
        return result

    start_p, end_p = range_info
    for i in range(start_p, end_p):
        if i >= len(doc.paragraphs):
            break
        text = doc.paragraphs[i].text.strip()
        if text.startswith("• 클렌징:"):
            result["cleansing"] = text[6:].strip()
        elif text.startswith("• 오일 사용법:"):
            result["oil"] = text[9:].strip()
        elif text.startswith("• 홈케어 추천:"):
            result["recommended"] = text[9:].strip()
        elif text.startswith("• 두피 열감:"):
            result["scalp"] = text[8:].strip()
        elif text.startswith("• 고농도"):
            result["peeling"] = text[2:].strip()

    return result


def extract_all_types(doc):
    """Main extraction function for all 16 types."""
    types_data = []

    for type_code in TYPE_CODES:
        print(f"  Extracting {type_code}...")

        # Get nickname and name from TIPY card table
        card_info, _ = extract_type_overview_card(doc, type_code)

        # Get characteristics and paragraph-based content
        para_content = extract_type_section_paragraphs(doc, type_code)

        # Get 3-step care from Ch05 STEP cards
        step_cards = extract_ch05_step_cards(doc, type_code)

        # Get management manual steps (Part 2 detailed)
        manual_steps = extract_management_manual_steps(doc, type_code)
        # Use manual steps if step cards unavailable or have "undefined"
        three_step = manual_steps if manual_steps else step_cards

        # Get in-care program table
        incare_program = extract_in_care_program(doc, type_code)

        # Get homecare from Part 2 table
        homecare_table = extract_homecare_table(doc, type_code)

        # Get homecare bullets from Ch05 paragraphs (for types that have them)
        homecare_bullets = extract_homecare_bullets(doc, type_code)
        # Merge: prefer table data for Part 2, use bullets as supplement
        homecare = {**homecare_bullets, **homecare_table}

        # Get hooking script from Part 2 table
        hooking_table = extract_hooking_script(doc, type_code)

        # Get diagnostic checklist
        checklist = extract_diagnostic_checklist(doc, type_code)

        # Build indicators map from type code letters
        code_chars = list(type_code)
        indicators = {
            "moisture": code_chars[0],   # O or D
            "sensitivity": code_chars[1],  # S or R
            "pigment": code_chars[2],    # N or P
            "aging": code_chars[3],      # T or W
        }

        type_entry = {
            "code": type_code,
            "nickname": card_info.get("nickname", ""),
            "name": card_info.get("name", ""),
            "indicators": indicators,
            "characteristics": para_content.get("characteristics", ""),
            "coreStrategy": para_content.get("coreStrategy", ""),
            "threeStepCare": three_step,
            "inCareProgram": incare_program,
            "homecare": homecare,
            "hookingScript": hooking_table or para_content.get("hookingScript", ""),
            "consultationPoints": para_content.get("consultationPoints", []),
            "diagnosticChecklist": checklist,
        }
        types_data.append(type_entry)

    return types_data


def js_escape(s):
    """Escape string for JS template literal."""
    if not isinstance(s, str):
        return ""
    return s.replace("\\", "\\\\").replace("`", "\\`").replace("${", "\\${")


def build_js_output(types_data, ch01, ch02, ch03, ch04):
    """Build the JS output file content."""
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    lines = []
    lines.append(f"// Auto-generated from 스킨_MBTI_책자_완성본 part 2.docx")
    lines.append(f"// Generated: {ts}")
    lines.append("")
    lines.append("const SKIN_TYPES_DATA = [")

    for i, t in enumerate(types_data):
        comma = "," if i < len(types_data) - 1 else ""

        # Build three step care
        tsc = t.get("threeStepCare", {})
        tsc_s1 = tsc.get("step1", {})
        tsc_s2 = tsc.get("step2", {})
        tsc_s3 = tsc.get("step3", {})

        # Build homecare
        hc = t.get("homecare", {})

        # Build in-care program as array
        icp = t.get("inCareProgram", [])
        icp_lines = []
        for j, item in enumerate(icp):
            stage = js_escape(item.get("stage", ""))
            content = js_escape(item.get("content", ""))
            tip = js_escape(item.get("tip", ""))
            item_comma = "," if j < len(icp) - 1 else ""
            icp_lines.append(f'        {{ stage: `{stage}`, content: `{content}`, tip: `{tip}` }}{item_comma}')

        # Build consultation points
        cp = t.get("consultationPoints", [])
        cp_lines = []
        for j, p in enumerate(cp):
            item_comma = "," if j < len(cp) - 1 else ""
            cp_lines.append(f'      `{js_escape(p)}`{item_comma}')

        # Build diagnostic checklist
        dc = t.get("diagnosticChecklist", {})
        dc_keys = list(dc.keys())
        dc_entries = []
        for j, (k, items) in enumerate(dc.items()):
            items_js = ", ".join(f'`{js_escape(item)}`' for item in items)
            entry_comma = "," if j < len(dc_keys) - 1 else ""
            dc_entries.append(f'      {k}: [{items_js}]{entry_comma}')

        ind = t.get("indicators", {})

        block = f"""  {{
    code: "{t['code']}",
    nickname: `{js_escape(t['nickname'])}`,
    name: `{js_escape(t['name'])}`,
    indicators: {{ moisture: "{ind.get('moisture','')}", sensitivity: "{ind.get('sensitivity','')}", pigment: "{ind.get('pigment','')}", aging: "{ind.get('aging','')}" }},
    characteristics: `{js_escape(t['characteristics'])}`,
    coreStrategy: `{js_escape(t['coreStrategy'])}`,
    threeStepCare: {{
      step1: {{ title: `{js_escape(tsc_s1.get('title',''))}`, description: `{js_escape(tsc_s1.get('description',''))}` }},
      step2: {{ title: `{js_escape(tsc_s2.get('title',''))}`, description: `{js_escape(tsc_s2.get('description',''))}` }},
      step3: {{ title: `{js_escape(tsc_s3.get('title',''))}`, description: `{js_escape(tsc_s3.get('description',''))}` }}
    }},
    inCareProgram: [
{chr(10).join(icp_lines)}
    ],
    homecare: {{
      cleansing: `{js_escape(hc.get('cleansing',''))}`,
      oil: `{js_escape(hc.get('oil',''))}`,
      moisturizing: `{js_escape(hc.get('moisturizing',''))}`,
      avoid: `{js_escape(hc.get('avoid',''))}`,
      special: `{js_escape(hc.get('special',''))}`,
      recommended: `{js_escape(hc.get('recommended',''))}`,
      scalp: `{js_escape(hc.get('scalp',''))}`
    }},
    hookingScript: `{js_escape(t['hookingScript'])}`,
    consultationPoints: [
{chr(10).join(cp_lines)}
    ],
    diagnosticChecklist: {{
{chr(10).join(dc_entries)}
    }}
  }}{comma}"""
        lines.append(block)

    lines.append("];")
    lines.append("")

    # CHAPTERS_DATA
    lines.append("const CHAPTERS_DATA = {")

    # ch01
    lines.append("  ch01_indicators: [")
    for i, ind in enumerate(ch01):
        comma = "," if i < len(ch01) - 1 else ""
        left = ind["left"]
        right = ind["right"]
        lines.append(f"""    {{
      pair: `{js_escape(ind['pair'])}`,
      title: `{js_escape(ind['title'])}`,
      subtitle: `{js_escape(ind['subtitle'])}`,
      left: {{ code: "{left['code']}", name: `{js_escape(left['name'])}`, description: `{js_escape(left['description'])}` }},
      right: {{ code: "{right['code']}", name: `{js_escape(right['name'])}`, description: `{js_escape(right['description'])}` }}
    }}{comma}""")
    lines.append("  ],")
    lines.append("")

    # ch02
    lines.append("  ch02_solutions: [")
    for i, sol in enumerate(ch02):
        comma = "," if i < len(ch02) - 1 else ""
        lines.append(f"""    {{
      indicator: `{js_escape(sol['indicator'])}`,
      goal: `{js_escape(sol['goal'])}`,
      goalDescription: `{js_escape(sol['goalDescription'])}`,
      equipment: `{js_escape(sol['equipment'])}`,
      tip: `{js_escape(sol['tip'])}`
    }}{comma}""")
    lines.append("  ],")
    lines.append("")

    # ch03
    lines.append("  ch03_overview: [")
    for i, card in enumerate(ch03):
        comma = "," if i < len(ch03) - 1 else ""
        lines.append(f"""    {{
      code: "{card['code']}",
      nickname: `{js_escape(card['nickname'])}`,
      summary: `{js_escape(card['summary'])}`,
      arrow: `{js_escape(card['arrow'])}`
    }}{comma}""")
    lines.append("  ],")
    lines.append("")

    # ch04
    s1 = ch04.get("step1", {})
    s2 = ch04.get("step2", {})
    s3 = ch04.get("step3", {})
    lines.append("  ch04_roadmap: {")
    for step_key, step_data in [("step1", s1), ("step2", s2), ("step3", s3)]:
        comma = "," if step_key != "step3" else ""
        lines.append(f"""    {step_key}: {{
      title: `{js_escape(step_data.get('title',''))}`,
      subtitle: `{js_escape(step_data.get('subtitle',''))}`,
      goal: `{js_escape(step_data.get('goal',''))}`,
      equipment: `{js_escape(step_data.get('equipment',''))}`,
      customerExperience: `{js_escape(step_data.get('customerExperience',''))}`
    }}{comma}""")
    lines.append("  }")
    lines.append("};")
    lines.append("")

    return "\n".join(lines)


def main():
    print(f"Loading docx: {DOCX_PATH}")
    doc = docx.Document(DOCX_PATH)
    print(f"  Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")

    print("\nExtracting chapter data...")
    ch01 = extract_ch01_indicators(doc)
    print(f"  Ch01 indicators: {len(ch01)}")

    ch02 = extract_ch02_solutions(doc)
    print(f"  Ch02 solutions: {len(ch02)}")

    ch03 = extract_ch03_overview(doc)
    print(f"  Ch03 overview cards: {len(ch03)}")

    ch04 = extract_ch04_roadmap(doc)
    print(f"  Ch04 roadmap steps: {len(ch04)}")

    print("\nExtracting 16 type data...")
    types_data = extract_all_types(doc)

    # Summary
    print(f"\n=== Extraction Summary ===")
    print(f"Types extracted: {len(types_data)}")
    for t in types_data:
        code = t["code"]
        nickname = t["nickname"]
        checklist_counts = {k: len(v) for k, v in t["diagnosticChecklist"].items()}
        consultation_count = len(t["consultationPoints"])
        incare_count = len(t["inCareProgram"])
        print(f"  {code} {nickname}: checklist={checklist_counts}, consultation={consultation_count}, incare={incare_count}")

    # Build JS output
    print(f"\nBuilding JS output...")
    js_content = build_js_output(types_data, ch01, ch02, ch03, ch04)

    # Write output
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    with open(OUTPUT_PATH, "w", encoding="utf-8") as f:
        f.write(js_content)

    print(f"\nOutput written to: {OUTPUT_PATH}")
    print(f"File size: {os.path.getsize(OUTPUT_PATH):,} bytes")

    # Verify output
    print("\n=== Verification ===")
    with open(OUTPUT_PATH, "r", encoding="utf-8") as f:
        content = f.read()
    code_check = [code for code in TYPE_CODES if f'code: "{code}"' in content]
    print(f"Type codes found in output: {len(code_check)}/16 -> {code_check}")
    print(f"SKIN_TYPES_DATA defined: {'const SKIN_TYPES_DATA' in content}")
    print(f"CHAPTERS_DATA defined: {'const CHAPTERS_DATA' in content}")
    print(f"ch01_indicators: {'ch01_indicators' in content}")
    print(f"ch02_solutions: {'ch02_solutions' in content}")
    print(f"ch03_overview: {'ch03_overview' in content}")
    print(f"ch04_roadmap: {'ch04_roadmap' in content}")


if __name__ == "__main__":
    main()
